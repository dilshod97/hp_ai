"""Chat suhbatini PDF yoki DOCX formatga eksport.

PDF — ReportLab
DOCX — python-docx
"""
import io
import os
import re
from datetime import datetime

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor

from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.lib import colors

# DejaVu — kirill/lotin/oʻzbek belgilarini qoʻllab-quvvatlaydi
DEJAVU_REGULAR = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
DEJAVU_BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"

_FONTS_REGISTERED = False


def _register_fonts() -> None:
    global _FONTS_REGISTERED
    if _FONTS_REGISTERED:
        return
    try:
        if os.path.exists(DEJAVU_REGULAR):
            pdfmetrics.registerFont(TTFont("UZ", DEJAVU_REGULAR))
        if os.path.exists(DEJAVU_BOLD):
            pdfmetrics.registerFont(TTFont("UZ-Bold", DEJAVU_BOLD))
        _FONTS_REGISTERED = True
    except Exception:
        pass


def _strip_md(text: str) -> str:
    """Markdownni oddiy matnga aylantirish (eksport uchun)."""
    if not text:
        return ""
    t = text
    # Bold/italic
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", t)
    t = re.sub(r"\*(.+?)\*", r"\1", t)
    # Inline code
    t = re.sub(r"`([^`]+)`", r"\1", t)
    # Code blocks
    t = re.sub(r"```[a-z]*\n(.*?)```", r"\1", t, flags=re.DOTALL)
    # Headers — # ni olib tashlash
    t = re.sub(r"^#{1,6}\s+", "", t, flags=re.MULTILINE)
    # Links [text](url) -> text
    t = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", t)
    return t.strip()


def _fmt_dt(iso: str | None) -> str:
    if not iso:
        return ""
    try:
        dt = datetime.fromisoformat(iso)
        return dt.strftime("%Y-%m-%d %H:%M")
    except Exception:
        return iso


# ---------- PDF ----------

def export_pdf(
    workspace: dict,
    messages: list[dict],
    user: dict | None = None,
) -> bytes:
    """Chat tarixini PDF qilib qaytaradi (bytes)."""
    _register_fonts()
    font = "UZ" if _FONTS_REGISTERED else "Helvetica"
    font_bold = "UZ-Bold" if _FONTS_REGISTERED else "Helvetica-Bold"

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=2 * cm, leftMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
        title=f"HP-AI audit assistant — {workspace.get('name', 'Chat')}",
        author="HP-AI audit assistant",
        subject="Audit chat tarixi",
    )
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "title", parent=styles["Heading1"], fontName=font_bold,
        fontSize=18, leading=22, textColor=colors.HexColor("#1e40af"),
        spaceAfter=6,
    )
    sub_style = ParagraphStyle(
        "sub", parent=styles["Normal"], fontName=font,
        fontSize=10, leading=14, textColor=colors.HexColor("#6b7280"),
        spaceAfter=12,
    )
    user_style = ParagraphStyle(
        "user", parent=styles["Normal"], fontName=font_bold,
        fontSize=11, leading=15, textColor=colors.HexColor("#1e3a8a"),
        spaceAfter=4, spaceBefore=10,
    )
    bot_label_style = ParagraphStyle(
        "bot_label", parent=styles["Normal"], fontName=font_bold,
        fontSize=11, leading=15, textColor=colors.HexColor("#15803d"),
        spaceAfter=4, spaceBefore=8,
    )
    body_style = ParagraphStyle(
        "body", parent=styles["Normal"], fontName=font,
        fontSize=10, leading=14, alignment=TA_LEFT, spaceAfter=6,
    )
    meta_style = ParagraphStyle(
        "meta", parent=styles["Normal"], fontName=font,
        fontSize=8, leading=10, textColor=colors.HexColor("#9ca3af"),
        spaceAfter=4,
    )

    elements = []

    # Sarlavha
    elements.append(Paragraph(
        f"{workspace.get('icon', '💬')} {workspace.get('name', 'Chat')}",
        title_style,
    ))

    # Meta sarlavha
    meta_lines = []
    if user:
        meta_lines.append(f"Foydalanuvchi: {user.get('full_name') or user.get('username')}")
        if user.get("sector"):
            meta_lines.append(f"Soha: {user['sector']}")
    meta_lines.append(f"Eksport sanasi: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    meta_lines.append(f"Xabarlar soni: {len(messages)}")
    if workspace.get("system_prompt"):
        meta_lines.append(f"Rol: {workspace['system_prompt'][:200]}")
    elements.append(Paragraph(" • ".join(meta_lines), sub_style))

    # Ajratuvchi chiziq
    elements.append(Table(
        [[""]], colWidths=[doc.width],
        style=TableStyle([
            ("LINEBELOW", (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
        ]),
    ))
    elements.append(Spacer(1, 12))

    # Xabarlar
    for msg in messages:
        role = msg.get("role")
        text = _strip_md(msg.get("text", "")).replace("\n", "<br/>")
        # XSS uchun emas — ReportLab Paragraph HTML belgilari uchun tozalash kerak
        text = (text
                .replace("&", "&amp;")
                .replace("<br/>", "BREAK_HERE")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace("BREAK_HERE", "<br/>"))
        ts = _fmt_dt(msg.get("created_at"))

        if role == "user":
            elements.append(Paragraph("👤 Savol", user_style))
            elements.append(Paragraph(text, body_style))
            if ts:
                elements.append(Paragraph(ts, meta_style))
        elif role == "bot":
            elements.append(Paragraph("🤖 Javob", bot_label_style))
            elements.append(Paragraph(text, body_style))
            # Vaqt va tezlik
            meta_parts = []
            if ts:
                meta_parts.append(ts)
            if msg.get("elapsed_ms"):
                meta_parts.append(f"{msg['elapsed_ms']/1000:.1f}s")
            if msg.get("tokens_per_sec"):
                meta_parts.append(f"{msg['tokens_per_sec']:.1f} tok/s")
            if meta_parts:
                elements.append(Paragraph(" • ".join(meta_parts), meta_style))
        elif role == "attach":
            elements.append(Paragraph(f"📎 {text}", meta_style))

    doc.build(elements)
    return buf.getvalue()


# ---------- DOCX ----------

def export_docx(
    workspace: dict,
    messages: list[dict],
    user: dict | None = None,
) -> bytes:
    """Chat tarixini Word (DOCX) qilib qaytaradi (bytes)."""
    doc = DocxDocument()

    # Default font (DejaVu mavjud bo'lsa)
    style = doc.styles["Normal"]
    style.font.name = "DejaVu Sans"
    style.font.size = Pt(11)

    # Sarlavha
    title = doc.add_heading(
        f"{workspace.get('icon', '')} {workspace.get('name', 'Chat')}".strip(),
        level=1,
    )
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    # Meta
    meta_p = doc.add_paragraph()
    meta_lines = []
    if user:
        meta_lines.append(f"Foydalanuvchi: {user.get('full_name') or user.get('username')}")
        if user.get("sector"):
            meta_lines.append(f"Soha: {user['sector']}")
    meta_lines.append(f"Eksport: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    meta_lines.append(f"Xabarlar: {len(messages)}")
    meta_run = meta_p.add_run(" • ".join(meta_lines))
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    if workspace.get("system_prompt"):
        rp = doc.add_paragraph()
        run = rp.add_run(f"Rol: {workspace['system_prompt']}")
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()  # bo'sh qator

    # Xabarlar
    for msg in messages:
        role = msg.get("role")
        text = _strip_md(msg.get("text", ""))
        ts = _fmt_dt(msg.get("created_at"))

        if role == "user":
            p = doc.add_paragraph()
            run = p.add_run("👤 Savol")
            run.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            run.font.size = Pt(11)

            doc.add_paragraph(text)

            if ts:
                meta = doc.add_paragraph()
                mr = meta.add_run(ts)
                mr.font.size = Pt(8)
                mr.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

        elif role == "bot":
            p = doc.add_paragraph()
            run = p.add_run("🤖 Javob")
            run.bold = True
            run.font.color.rgb = RGBColor(0x15, 0x80, 0x3D)
            run.font.size = Pt(11)

            doc.add_paragraph(text)

            meta_parts = []
            if ts:
                meta_parts.append(ts)
            if msg.get("elapsed_ms"):
                meta_parts.append(f"{msg['elapsed_ms']/1000:.1f}s")
            if msg.get("tokens_per_sec"):
                meta_parts.append(f"{msg['tokens_per_sec']:.1f} tok/s")
            if meta_parts:
                meta = doc.add_paragraph()
                mr = meta.add_run(" • ".join(meta_parts))
                mr.font.size = Pt(8)
                mr.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

        elif role == "attach":
            p = doc.add_paragraph()
            run = p.add_run(f"📎 {text}")
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
