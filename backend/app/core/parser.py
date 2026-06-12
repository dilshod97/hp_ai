"""Fayl turlarini matnga aylantirish va chunklarga boʻlish.

Qoʻllab-quvvatlanadi: PDF (matn + OCR), DOCX, XLSX, TXT.
"""
import logging
import os
import re
from typing import Iterable

import pdfplumber
from docx import Document as DocxDocument
from openpyxl import load_workbook

from app.config import settings

log = logging.getLogger("parser")

# OCR til kombinatsiyasi — Tesseract format ("lang1+lang2+...")
# uzb (lotin), uzb_cyrl (kiril), rus, eng — barchasini birga sinaymiz
OCR_LANG = "uzb+uzb_cyrl+rus+eng"
OCR_MIN_TEXT_PER_PAGE = 50  # belgi — bundan kam bo'lsa, OCR ishga tushadi
OCR_DPI = 200  # ko'p — sifat, kam — tezlik. 200 — yaxshi muvozanat


def _ocr_pdf_page(path: str, page_num: int) -> str:
    """Bitta PDF sahifani Tesseract bilan OCR qiladi."""
    try:
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError as e:
        log.warning("OCR uchun kutubxonalar yo'q: %s", e)
        return ""

    try:
        images = convert_from_path(
            path, dpi=OCR_DPI, first_page=page_num, last_page=page_num
        )
        if not images:
            return ""
        text = pytesseract.image_to_string(images[0], lang=OCR_LANG)
        return text.strip()
    except Exception as e:
        log.warning("OCR xatosi (sahifa %d): %s", page_num, e)
        return ""


def _read_pdf(path: str) -> str:
    """PDF'dan matn oladi. Agar sahifada matn yo'q bo'lsa (skaner) — OCR ishlaydi."""
    parts: list[str] = []
    ocr_pages = 0

    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()

            # Agar sahifa "bo'sh" (skaner bo'lishi mumkin) — OCR ishga tushiramiz
            if len(text) < OCR_MIN_TEXT_PER_PAGE:
                ocr_text = _ocr_pdf_page(path, i)
                if ocr_text and len(ocr_text) > len(text):
                    text = ocr_text
                    ocr_pages += 1

            parts.append(text)

            # Jadvallarni ham qoʻshish
            for table in page.extract_tables() or []:
                for row in table:
                    if not row:
                        continue
                    parts.append(" | ".join(str(c) if c else "" for c in row))

    if ocr_pages:
        log.info("📷 OCR ishlatildi: %d sahifa (%s)", ocr_pages, os.path.basename(path))
    return "\n".join(parts)


def _read_docx(path: str) -> str:
    doc = DocxDocument(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_xlsx(path: str) -> str:
    wb = load_workbook(path, data_only=True, read_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"# Varaq: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _read_image(path: str) -> str:
    """Rasmga OCR qo'llab matn olish."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return ""
    try:
        img = Image.open(path)
        return pytesseract.image_to_string(img, lang=OCR_LANG).strip()
    except Exception as e:
        log.warning("Rasm OCR xato (%s): %s", path, e)
        return ""


def extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)
    if ext == ".xlsx":
        return _read_xlsx(path)
    if ext in (".txt", ".md"):
        return _read_txt(path)
    if ext in (".jpg", ".jpeg", ".png", ".tiff", ".bmp"):
        return _read_image(path)
    raise ValueError(f"Qoʻllab-quvvatlanmaydi: {ext}")


# ---------- Chunking ----------

_SPLIT_RE = re.compile(r"(?<=[\.\?\!])\s+|\n{2,}")


def _split_sentences(text: str) -> list[str]:
    parts = _SPLIT_RE.split(text)
    return [p.strip() for p in parts if p and p.strip()]


def chunk_text(
    text: str,
    chunk_size: int = None,
    overlap: int = None,
) -> Iterable[str]:
    """Soʻzlar boʻyicha chunking, gap chegarasini saqlashga harakat qiladi."""
    chunk_size = chunk_size or settings.CHUNK_SIZE
    overlap = overlap or settings.CHUNK_OVERLAP

    sentences = _split_sentences(text)
    if not sentences:
        return []

    chunks: list[str] = []
    buf: list[str] = []
    buf_len = 0

    for sent in sentences:
        w = sent.split()
        if buf_len + len(w) > chunk_size and buf:
            chunks.append(" ".join(buf))
            # overlap
            if overlap > 0:
                tail_words = " ".join(buf).split()[-overlap:]
                buf = [" ".join(tail_words)]
                buf_len = len(tail_words)
            else:
                buf, buf_len = [], 0
        buf.append(sent)
        buf_len += len(w)

    if buf:
        chunks.append(" ".join(buf))
    return chunks
